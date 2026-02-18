"""
DOCX document processor implementation.
Handles extraction of text and page mapping from DOCX files using python-docx.
"""

from typing import List, Tuple
from pathlib import Path
from docx import Document

from .base_processor import DocumentProcessor


class DOCXProcessor(DocumentProcessor):
    """
    Concrete implementation for processing DOCX files.

    Uses python-docx library to extract text from paragraphs and tables.
    Page numbers are estimated based on character count since DOCX doesn't
    have fixed pages like PDF.
    """

    def get_supported_extensions(self) -> List[str]:
        """DOCX and DOC files support."""
        return ['.docx', '.DOCX', '.doc', '.DOC']

    def validate_file(self, file_path: str) -> bool:
        """
        Validate if the file is a valid DOCX.

        Args:
            file_path: Path to the DOCX file

        Returns:
            True if file exists, has .docx extension, and can be opened
        """
        path = Path(file_path)

        # Check file exists
        if not path.exists():
            print(f"File does not exist: {file_path}")
            return False

        # Check extension
        if path.suffix.lower() not in ['.docx', '.doc']:
            print(f"File is not a DOCX: {file_path}")
            return False

        # Try to open and validate it's a valid DOCX
        try:
            doc = Document(file_path)
            # If we can create a Document object, it's valid
            return True
        except Exception as e:
            print(f"Error validating DOCX {file_path}: {e}")
            return False

    def extract_text(self, file_path: str) -> Tuple[str, List[Tuple[int, int, int]]]:
        """
        Extract text from DOCX file with estimated page tracking.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (full_text, page_mapping)
            where page_mapping is List[(start_pos, end_pos, estimated_page_num)]

        Raises:
            ValueError: If DOCX cannot be read
        """
        try:
            doc = Document(file_path)
            text = ""
            page_mapping = []
            chars_per_page = 2500  # Rough estimate for page breaks

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                start_pos = len(text)
                paragraph_text = paragraph.text + "\n"
                text += paragraph_text
                end_pos = len(text) - 1

                # Only track non-empty paragraphs
                if paragraph.text.strip():
                    # Estimate page number based on character position
                    estimated_page = max(1, (start_pos // chars_per_page) + 1)
                    page_mapping.append((start_pos, end_pos, estimated_page))

            # Extract text from tables
            for table in doc.tables:
                start_pos = len(text)
                table_text = ""
                for row in table.rows:
                    for cell in row.cells:
                        table_text += cell.text + " "
                    table_text += "\n"

                text += table_text
                end_pos = len(text) - 1

                if table_text.strip():
                    estimated_page = max(1, (start_pos // chars_per_page) + 1)
                    page_mapping.append((start_pos, end_pos, estimated_page))

            return text, page_mapping

        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX {file_path}: {e}")
